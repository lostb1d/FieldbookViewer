import sys
import os
import re
import sqlite3
import json
import io
import subprocess
import tempfile
from datetime import datetime

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGroupBox, QFormLayout,
    QLabel, QLineEdit, QPushButton, QComboBox, QFileDialog, QGraphicsView, QGraphicsScene,
    QGraphicsPixmapItem, QMessageBox, QListWidget, QStackedWidget, QDialog, QScrollArea, QSlider, QAction
)

from PyQt5.QtGui import (
    QPixmap, QIntValidator, QIcon, QPalette, QPainter, QPen, QImage, QClipboard, QTransform
)
from PyQt5.QtCore import Qt, QRectF, QPoint, QBuffer

from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import fitz  # PyMuPDF

def get_appdata_folder(appname="FieldbookViewer"):
    if sys.platform == "win32":
        # e.g., C:\Users\user\AppData\Local\FieldbookViewer
        base = os.environ.get('LOCALAPPDATA') or os.path.expanduser('~\\AppData\\Local')
        return os.path.join(base, appname)
    elif sys.platform == 'darwin':
        # e.g., /Users/user/Library/Application Support/FieldbookViewer
        return os.path.join(os.path.expanduser('~/Library/Application Support/'), appname)
    else:
        # Linux/UNIX: ~/.local/share/FieldbookViewer
        return os.path.join(os.path.expanduser('~/.local/share/'), appname)

def to_nepali_number(num):
    num_map = str.maketrans('0123456789', '०१२३४५६७८९')
    return str(num).translate(num_map)

class Config:
    def __init__(self, path):
        self.path = path
        self.data = {}
        self.load()
    def load(self):
        if os.path.exists(self.path):
            with open(self.path, "r", encoding="utf-8") as f:
                self.data = json.load(f)
        else:
            self.data = {}
    def save(self):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, indent=2)
    def get_folder(self, key):
        return self.data.get(key, "")
    def set_folder(self, key, folder):
        self.data[key] = folder
        self.save()

class UserDB:
    def __init__(self, db_path):
        self.conn = sqlite3.connect(db_path)
        self.create_table()
    def create_table(self):
        self.conn.execute('''CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL,
            role TEXT NOT NULL
        )''')
        cur = self.conn.cursor()
        cur.execute("SELECT * FROM users WHERE username='admin'")
        if not cur.fetchone():
            self.conn.execute("INSERT INTO users VALUES (?, ?, ?)", ('admin', 'admin', 'admin'))
            self.conn.commit()
    def validate(self, username, password):
        cur = self.conn.cursor()
        cur.execute("SELECT role FROM users WHERE username=? AND password=?", (username, password))
        row = cur.fetchone()
        return row[0] if row else None

class FieldbookBottomTextDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Fieldbook/Footer Details")
        self.inputs = {}
        form = QFormLayout()
        self.inputs["patra_pathaune"] = QLineEdit()
        form.addRow("पत्र पठाउने :", self.inputs["patra_pathaune"])
        self.inputs["chan_dan"] = QLineEdit()
        form.addRow("च.नं./द.नं. :", self.inputs["chan_dan"])
        self.inputs["miti"] = QLineEdit()
        form.addRow("मिति :", self.inputs["miti"])
        self.inputs["prayojan"] = QLineEdit()
        form.addRow("प्रयोजन :", self.inputs["prayojan"])
        self.inputs["rasid_no"] = QLineEdit()
        form.addRow("रसिद नं :", self.inputs["rasid_no"])
        button = QPushButton("OK")
        button.clicked.connect(self.accept)
        layout = QVBoxLayout(self)
        layout.addLayout(form)
        layout.addWidget(button)
    def get_values(self):
        return {k: field.text().strip() for k, field in self.inputs.items()}

class FieldbookDocManager:
    def __init__(self):
        self.doc = None
        self.section = None
        self.images_on_page = 0
        self.max_images_per_page = 3
        self.loaded_template = None
        self.footer_info = None
    def new_from_template(self, template_path):
        self.doc = Document(template_path)
        self.loaded_template = template_path
        self.section = self.doc.sections[0]
        self.images_on_page = 0
        self.footer_info = None
    def get_footer_line(self):
        info = self.footer_info or {}
        def safe(k, dots):
            val = info.get(k, "")
            return val if val else dots
        return (
            f"श्री {safe('patra_pathaune','....................')} को च.नं./द.नं. {safe('chan_dan','.......')} मिति {safe('miti','.............')} "
            f"को पत्रानुसार {safe('prayojan','...............')} प्रयोजनको लागि  "
            f"रसिद नं {safe('rasid_no','.....................')} बाट राजश्व लिई कम्प्युटरबाट फिल्डबुक/प्लट रजिष्टर प्रतिलिपि उतार गरि पठाइएको व्यहोरा अनुरोध छ ।"
        )
    def insert_footer_to_all_pages(self, footer_info):
        self.footer_info = footer_info
        section = self.doc.sections[0]
        footer = section.footer
        for element in footer._element.xpath("./w:p | ./w:tbl"):
            footer._element.remove(element)
        p1 = footer.add_paragraph()
        run1 = p1.add_run(self.get_footer_line())
        run1.font.size = Pt(10)
        run1.font.name = "Kalimati"
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kalimati')
        p1.alignment = 1
        table = footer.add_table(rows=1, cols=3, width=section.page_width - section.left_margin - section.right_margin)
        table.allow_autofit = True
        cell_texts = ["प्रिन्ट गर्ने", "रुजु गर्ने", "प्रमाणित गर्ने"]
        aligns = [0, 1, 2]
        for i, text in enumerate(cell_texts):
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Kalimati"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kalimati')
            p.alignment = aligns[i]
        tbl = table._tbl
        for cell in tbl.iter():
            if cell.tag.endswith('tcBorders'):
                cell.getparent().remove(cell)
        section.footer_distance = Pt(10)
    def add_image(self, pil_img, vdc, ward, sheet, parcel):
        meta_text = (
            f"गा.वि.स: {vdc} | वडा नं: {to_nepali_number(ward)} | सिट: {to_nepali_number(sheet)} | कित्ता नं: {to_nepali_number(parcel)}"
        )
        avail_width = self.section.page_width - self.section.left_margin - self.section.right_margin
        temp_io = io.BytesIO()
        pil_img.save(temp_io, format="PNG")
        temp_io.seek(0)
        if self.images_on_page >= self.max_images_per_page:
            self.doc.add_page_break()
            self.images_on_page = 0
        p = self.doc.add_paragraph(meta_text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Kalimati"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kalimati')
        self.doc.add_picture(temp_io, width=avail_width)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.paragraph_format.space_after = Pt(0)
        last_paragraph.paragraph_format.keep_with_next = True
        self.images_on_page += 1
    def save(self, path):
        if self.footer_info is not None:
            self.insert_footer_to_all_pages(self.footer_info)
        if self.doc:
            self.doc.save(path)
    def is_loaded(self):
        return self.doc is not None
    def close(self):
        self.doc = None
        self.section = None
        self.loaded_template = None
        self.images_on_page = 0
        self.footer_info = None

fieldbook_doc_mgr = FieldbookDocManager()
plotregister_doc_mgr = FieldbookDocManager()

class PDFPreviewDialog(QDialog):
    def __init__(self, pdf_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Print Preview")
        layout = QVBoxLayout(self)
        self.scroll_area = QScrollArea()
        widget = QWidget()
        vbox = QVBoxLayout(widget)
        doc = fitz.open(pdf_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=120)
            img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format_RGBA8888)
            label = QLabel()
            label.setPixmap(QPixmap.fromImage(img))
            vbox.addWidget(label)
        widget.setLayout(vbox)
        self.scroll_area.setWidget(widget)
        self.scroll_area.setWidgetResizable(True)
        layout.addWidget(self.scroll_area)
        btn_layout = QHBoxLayout()
        self.print_btn = QPushButton("Print")
        self.exit_btn = QPushButton("Exit")
        btn_layout.addWidget(self.print_btn)
        btn_layout.addWidget(self.exit_btn)
        layout.addLayout(btn_layout)
        self.exit_btn.clicked.connect(self.reject)
        self.print_btn.clicked.connect(lambda: self.print_pdf(pdf_path))
    def print_pdf(self, pdf_path):
        import platform
        try:
            if platform.system() == "Windows":
                os.startfile(pdf_path, "print")
            elif platform.system() == "Darwin":
                subprocess.run(["open", "-a", "Preview", pdf_path])
            else:
                subprocess.run(["lp", pdf_path])
        except Exception as e:
            QMessageBox.warning(self, "Print Error", f"Could not print: {str(e)}")

def convert_docx_to_pdf(docx_path, pdf_path):
    if sys.platform.startswith('win'):
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    else:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', os.path.dirname(pdf_path)
        ], check=True)
        pdf_generated = os.path.join(os.path.dirname(pdf_path), os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if pdf_generated != pdf_path and os.path.exists(pdf_generated):
            os.rename(pdf_generated, pdf_path)
        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF was not generated. Check if LibreOffice is installed and in PATH.")

class EnhancedImageViewer(QGraphicsView):
    def __init__(self, image_path=None):
        super().__init__()
        self.setScene(QGraphicsScene())
        self.base_pixmap = QPixmap(image_path) if image_path else QPixmap()
        self.angle = 0
        self.pixmap_item = QGraphicsPixmapItem(self.base_pixmap)
        self.scene().addItem(self.pixmap_item)
        self.setRenderHint(QPainter.Antialiasing)
        self.setRenderHint(QPainter.SmoothPixmapTransform)
        self.setTransformationAnchor(QGraphicsView.AnchorUnderMouse)
        self.setDragMode(QGraphicsView.NoDrag)
        self._zoom = 1.0
        self._pan = False
        self._pan_start = QPoint()
    def load_image(self, path):
        self.scene().clear()
        self.base_pixmap = QPixmap(path)
        self.angle = 0
        self.pixmap_item = QGraphicsPixmapItem(self.base_pixmap)
        self.scene().addItem(self.pixmap_item)
        self.setSceneRect(QRectF(self.base_pixmap.rect()))
        self.resetTransform()
        self._zoom = 1.0
    def set_rotation(self, angle):
        self.angle = angle
        t = QTransform()
        t.rotate(self.angle)
        self.pixmap_item.setPixmap(self.base_pixmap.transformed(t, Qt.SmoothTransformation))
    def wheelEvent(self, event):
        zoom_factor = 1.25 if event.angleDelta().y() > 0 else 0.8
        self.scale(zoom_factor, zoom_factor)
        self._zoom *= zoom_factor
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._pan = True
            self.setCursor(Qt.ClosedHandCursor)
            self._pan_start = event.pos()
        super().mousePressEvent(event)
    def mouseMoveEvent(self, event):
        if self._pan and event.buttons() & Qt.LeftButton:
            delta = self._pan_start - event.pos()
            self._pan_start = event.pos()
            self.horizontalScrollBar().setValue(self.horizontalScrollBar().value() + delta.x())
            self.verticalScrollBar().setValue(self.verticalScrollBar().value() + delta.y())
        super().mouseMoveEvent(event)
    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._pan = False
            self.setCursor(Qt.ArrowCursor)
        super().mouseReleaseEvent(event)
    def zoom_in(self):
        self._zoom *= 1.25
        self.scale(1.25, 1.25)
    def zoom_out(self):
        self._zoom *= 0.8
        self.scale(0.8, 0.8)
    def reset_view(self):
        self.resetTransform()
        self._zoom = 1.0

class ImageViewerWindow(QMainWindow):
    def __init__(self, image_path, config=None, meta=None, doc_type="fieldbook"):
        super().__init__()
        self.setWindowTitle("Image Viewer")
        self.viewer = EnhancedImageViewer(image_path)
        self._crop_mode = False
        self._last_crop = None
        self._rect_item = None
        self.config = config
        self.meta = meta or {}
        self.doc_type = doc_type
        self.rotation_slider = QSlider(Qt.Horizontal)
        self.rotation_slider.setMinimum(0)
        self.rotation_slider.setMaximum(360)
        self.rotation_slider.setValue(0)
        self.rotation_slider.setTickPosition(QSlider.TicksBelow)
        self.rotation_slider.setTickInterval(30)
        self.rotation_slider.valueChanged.connect(self.on_slider_rotate)
        self.meta_label = QLabel(self.format_metadata())
        self.meta_label.setWordWrap(True)
        self.meta_label.setStyleSheet("font-size: 15px; padding: 7px 3px; font-weight: 600; color: #192a60")
        btn_zoom_in = QPushButton("Zoom In")
        btn_zoom_out = QPushButton("Zoom Out")
        btn_crop = QPushButton("Crop")
        btn_copy = QPushButton("Copy Crop")
        btn_paste_to_word = QPushButton(f"Paste into {'Fieldbook' if doc_type == 'fieldbook' else 'Plot Register'} Word")
        btn_preview_print = QPushButton("Print Preview")
        btn_zoom_in.clicked.connect(self.viewer.zoom_in)
        btn_zoom_out.clicked.connect(self.viewer.zoom_out)
        btn_crop.clicked.connect(self.activate_crop)
        btn_copy.clicked.connect(self.copy_crop)
        btn_paste_to_word.clicked.connect(self.paste_to_word)
        btn_preview_print.clicked.connect(self.preview_print)
        btn_layout = QHBoxLayout()
        btn_layout.addWidget(btn_zoom_in)
        btn_layout.addWidget(btn_zoom_out)
        btn_layout.addWidget(btn_crop)
        btn_layout.addWidget(btn_copy)
        btn_layout.addWidget(btn_paste_to_word)
        btn_layout.addWidget(btn_preview_print)
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.meta_label)
        main_layout.addWidget(self.viewer)
        main_layout.addLayout(btn_layout)
        main_layout.addWidget(QLabel("Rotate (0°–360°):"))
        main_layout.addWidget(self.rotation_slider)
        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)
        self.viewer.viewport().installEventFilter(self)
    def on_slider_rotate(self, value):
        self.viewer.set_rotation(value)
    def format_metadata(self):
        m = self.meta
        return f"गा.वि.स: {m.get('vdc','')} | वडा नं: {m.get('ward','')} | सिट: {m.get('sheet','')} | कि. नं: {m.get('parcel','')}"
    def activate_crop(self):
        self._crop_mode = True
        self.viewer.setCursor(Qt.CrossCursor)
    def eventFilter(self, obj, event):
        if obj is self.viewer.viewport() and self._crop_mode:
            if event.type() == event.MouseButtonPress and event.button() == Qt.LeftButton:
                self._start = self.viewer.mapToScene(event.pos())
                if self._rect_item:
                    self.viewer.scene().removeItem(self._rect_item)
                    self._rect_item = None
                return True
            elif event.type() == event.MouseMove and hasattr(self, "_start") and self._start:
                end = self.viewer.mapToScene(event.pos())
                rect = QRectF(self._start, end).normalized()
                if self._rect_item:
                    self.viewer.scene().removeItem(self._rect_item)
                self._rect_item = self.viewer.scene().addRect(rect, pen=QPen(Qt.red, 2))
                return True
            elif event.type() == event.MouseButtonRelease and event.button() == Qt.LeftButton:
                self._crop_mode = False
                self.viewer.setCursor(Qt.ArrowCursor)
                if self._rect_item:
                    rect = self._rect_item.rect().toRect()
                    cropped = self.viewer.base_pixmap.copy(rect)
                    self._last_crop = cropped
                    self._rect_item = None
                return True
        return super().eventFilter(obj, event)
    def copy_crop(self):
        if self._last_crop:
            QApplication.clipboard().setPixmap(self._last_crop)
            QMessageBox.information(self, "Copied", "Cropped image copied to clipboard.")
        else:
            QMessageBox.warning(self, "No Crop", "No crop selected/cropped yet.")
    def get_pil_image(self):
        if self._last_crop:
            qimg = self._last_crop.toImage()
        else:
            qimg = self.viewer.base_pixmap.toImage()
        if qimg.isNull():
            return None
        buf = QBuffer()
        buf.open(QBuffer.ReadWrite)
        qimg.save(buf, "PNG")
        pil_img = Image.open(io.BytesIO(buf.data()))
        return pil_img

    def get_doc_mgr(self):
        return fieldbook_doc_mgr if self.doc_type == "fieldbook" else plotregister_doc_mgr
    def get_template_path(self):
        if not self.config:
            return None
        return self.config.get_folder("fieldbook_template") if self.doc_type == "fieldbook" else self.config.get_folder("plotregister_template")

    def paste_to_word(self):
        pil_img = self.get_pil_image()
        if not pil_img:
            QMessageBox.warning(self, "Error", "No image (or cropped image) to insert.")
            return
        doc_mgr = self.get_doc_mgr()
        template_path = self.get_template_path()
        if not template_path or not os.path.isfile(template_path):
            QMessageBox.warning(self, "Template", f"No {self.doc_type.title()} template loaded. Use File > Load {self.doc_type.title()} Template.")
            return
        if not doc_mgr.is_loaded() or doc_mgr.loaded_template != template_path:
            doc_mgr.new_from_template(template_path)
        vdc = self.meta.get("vdc", "")
        ward = self.meta.get("ward", "")
        sheet = self.meta.get("sheet", "")
        parcel = self.meta.get("parcel", "")
        doc_mgr.add_image(pil_img, vdc, ward, sheet, parcel)
        QMessageBox.information(self, "Image Added",
            f"Image added to {self.doc_type.title()}. You can finalize and save from the button below image list when you're done."
        )
    def preview_print(self):
        pil_img = self.get_pil_image()
        if not pil_img:
            QMessageBox.warning(self, "Error", "No image to preview.")
            return
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        qt_img = QImage.fromData(buf.getvalue())
        pixmap = QPixmap.fromImage(qt_img)
        label = QLabel()
        label.setPixmap(pixmap)
        label.setScaledContents(True)
        label.setMinimumSize(min(pixmap.width(), 800), min(pixmap.height(), 1000))
        dlg = QDialog(self)
        dlg.setWindowTitle("Print Preview")
        layout = QVBoxLayout(dlg)
        meta_label = QLabel(self.format_metadata())
        meta_label.setStyleSheet("font-size:15px; font-weight:600;")
        layout.addWidget(meta_label)
        scroll = QScrollArea()
        scroll.setWidget(label)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        dlg.resize(900, 1100)
        dlg.exec_()

class LoginWidget(QWidget):
    def __init__(self, db, on_login):
        super().__init__()
        self.db = db
        self.on_login = on_login
        self.init_ui()
    def init_ui(self):
        layout = QVBoxLayout()
        group = QGroupBox("Login")
        form = QFormLayout(group)
        self.user_edit = QLineEdit()
        self.pass_edit = QLineEdit()
        self.pass_edit.setEchoMode(QLineEdit.Password)
        form.addRow("Username:", self.user_edit)
        form.addRow("Password:", self.pass_edit)
        btn = QPushButton("Login")
        btn.setMinimumHeight(40)
        btn.clicked.connect(self.try_login)
        form.addRow(btn)
        group.setContentsMargins(10, 10, 10, 10)
        layout.addWidget(group)
        self.setLayout(layout)
    def try_login(self):
        username = self.user_edit.text()
        password = self.pass_edit.text()
        role = self.db.validate(username, password)
        if role:
            self.on_login(username, role)
        else:
            QMessageBox.warning(self, "Login Failed", "Invalid username or password.")

class BookViewer(QWidget):
    def __init__(self, config, config_key, title, doc_type, on_back=None):
        super().__init__()
        self.doc_type = doc_type
        self.on_back = on_back
        self.config = config
        self.config_key = config_key
        self.title = title
        self.folder = self.config.get_folder(self.config_key)
        self.init_ui()

    def get_doc_mgr(self):
        return fieldbook_doc_mgr if self.doc_type == "fieldbook" else plotregister_doc_mgr
    def get_template_path(self):
        return self.config.get_folder("fieldbook_template") if self.doc_type == "fieldbook" else self.config.get_folder("plotregister_template")

    def init_ui(self):
        main_layout = QHBoxLayout(self)
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        group = QGroupBox(self.title)
        form = QFormLayout(group)
        self.vdc_combo = QComboBox()
        self.ward_combo = QComboBox()
        self.sheet_combo = QComboBox()
        self.parcel_edit = QLineEdit()
        self.parcel_edit.setValidator(QIntValidator(1, 99999))
        self.search_btn = QPushButton("Search")
        self.search_btn.setMinimumHeight(36)
        self.search_btn.clicked.connect(self.search_image)
        form.addRow("VDC:", self.vdc_combo)
        form.addRow("Ward:", self.ward_combo)
        form.addRow("Sheet:", self.sheet_combo)
        form.addRow("Parcel No:", self.parcel_edit)
        form.addRow(self.search_btn)
        group.setContentsMargins(10, 10, 10, 10)
        left_layout.addWidget(group)
        self.image_list = QListWidget()
        left_layout.addWidget(QLabel("Available Images:"))
        left_layout.addWidget(self.image_list)
        self.finalize_btn = QPushButton("Save")
        self.print_btn = QPushButton("Print")
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)
        btn_row.addWidget(self.finalize_btn)
        btn_row.addWidget(self.print_btn)
        btn_row.addStretch()
        left_layout.addLayout(btn_row)
        left_layout.addStretch()
        left_widget.setMinimumWidth(350)
        left_widget.setMaximumWidth(500)
        self.back_btn = QPushButton("< Back")
        self.back_btn.clicked.connect(self.handle_back)
        left_layout.insertWidget(0, self.back_btn)
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        self.viewer = EnhancedImageViewer()
        btn_zoom_in = QPushButton("Zoom In")
        btn_zoom_out = QPushButton("Zoom Out")
        btn_zoom_in.clicked.connect(self.viewer.zoom_in)
        btn_zoom_out.clicked.connect(self.viewer.zoom_out)
        btn_layout = QHBoxLayout()
        btn_layout.addWidget(btn_zoom_in)
        btn_layout.addWidget(btn_zoom_out)
        right_layout.addWidget(QLabel("Preview:"))
        right_layout.addWidget(self.viewer)
        right_layout.addLayout(btn_layout)
        right_widget.setMinimumWidth(400)
        main_layout.addWidget(left_widget, 35)
        main_layout.addWidget(right_widget, 65)
        self.setLayout(main_layout)
        self.vdc_combo.currentTextChanged.connect(self.update_wards)
        self.ward_combo.currentTextChanged.connect(self.update_sheets)
        self.sheet_combo.currentTextChanged.connect(self.update_images)
        self.image_list.currentTextChanged.connect(self.load_selected_image)
        self.populate_vdcs()
        self.finalize_btn.clicked.connect(self.finalize_doc)
        self.print_btn.clicked.connect(self.print_doc)

        # Set the "Back" button small and red
        self.back_btn.setMinimumSize(64, 32)
        self.back_btn.setMaximumSize(72, 36)
        self.back_btn.setStyleSheet("""
            QPushButton {
                background-color: #e53935;
                color: white;
                font-size: 14px;
                border-radius: 7px;
                margin: 3px;
                padding: 3px 12px;
                min-width: 50px;
            }
            QPushButton:hover {
                background-color: #b71c1c;
            }
        """)

        # Set the "Search" button small and green
        self.search_btn.setMinimumSize(64, 40)
        self.search_btn.setMaximumSize(90, 46)
        self.search_btn.setStyleSheet("""
            QPushButton {
                background-color: #43a047;
                color: white;
                font-size: 14px;
                border-radius: 7px;
                margin: 3px;
                padding: 3px 14px;
                min-width: 50px;
            }
            QPushButton:hover {
                background-color: #1b5e20;
            }
        """)


    def handle_back(self):
        QApplication.clipboard().clear(mode=QClipboard.Clipboard)
        QApplication.clipboard().clear(mode=QClipboard.Selection)
        if callable(self.on_back):
            self.on_back()

    def set_folder(self, folder):
        self.folder = folder
        self.populate_vdcs()
    def populate_vdcs(self):
        self.vdc_combo.clear()
        if not self.folder or not os.path.isdir(self.folder):
            return
        vdcs = [d for d in os.listdir(self.folder) if os.path.isdir(os.path.join(self.folder, d))]
        self.vdc_combo.addItems(vdcs)
        if vdcs:
            self.vdc_combo.setCurrentIndex(0)
            self.update_wards(vdcs[0])
    def update_wards(self, vdc):
        self.ward_combo.clear()
        vdc_path = os.path.join(self.folder, vdc)
        if not os.path.isdir(vdc_path):
            return
        wards = [d for d in os.listdir(vdc_path) if os.path.isdir(os.path.join(vdc_path, d))]
        self.ward_combo.addItems(wards)
        direct_images = [f for f in os.listdir(vdc_path) if re.match(r"(\d+)-(\d+)\.jpe?g", f, re.IGNORECASE)]
        if direct_images:
            self.ward_combo.addItem("(No Sheet)")
        if self.ward_combo.count() > 0:
            self.ward_combo.setCurrentIndex(0)
            self.update_sheets(self.ward_combo.currentText())
    def update_sheets(self, ward):
        vdc = self.vdc_combo.currentText()
        vdc_path = os.path.join(self.folder, vdc)
        self.sheet_combo.clear()
        if ward == "(No Sheet)":
            self.update_images("(No Sheet)")
            return
        ward_path = os.path.join(vdc_path, ward)
        if not os.path.isdir(ward_path):
            return
        sheets = [d for d in os.listdir(ward_path) if os.path.isdir(os.path.join(ward_path, d))]
        self.sheet_combo.addItems(sheets)
        if sheets:
            self.sheet_combo.setCurrentIndex(0)
            self.update_images(sheets[0])
    def update_images(self, sheet):
        vdc = self.vdc_combo.currentText()
        ward = self.ward_combo.currentText()
        self.image_list.clear()
        if ward == "(No Sheet)" or sheet == "(No Sheet)":
            vdc_path = os.path.join(self.folder, vdc)
            images = [f for f in os.listdir(vdc_path) if re.match(r"(\d+)-(\d+)\.jpe?g", f, re.IGNORECASE)]
            self.image_list.addItems(images)
            if images:
                self.image_list.setCurrentRow(0)
            return
        sheet_path = os.path.join(self.folder, vdc, ward, sheet)
        if not os.path.isdir(sheet_path):
            return
        images = [f for f in os.listdir(sheet_path) if re.match(r"(\d+)-(\d+)\.jpe?g", f, re.IGNORECASE)]
        self.image_list.addItems(images)
        if images:
            self.image_list.setCurrentRow(0)
    def load_selected_image(self, filename):
        vdc = self.vdc_combo.currentText()
        ward = self.ward_combo.currentText()
        sheet = self.sheet_combo.currentText()
        if ward == "(No Sheet)" or sheet == "(No Sheet)":
            path = os.path.join(self.folder, vdc, filename)
        else:
            path = os.path.join(self.folder, vdc, ward, sheet, filename)
        if os.path.isfile(path):
            self.viewer.load_image(path)
    def search_image(self):
        vdc = self.vdc_combo.currentText()
        ward = self.ward_combo.currentText()
        sheet = self.sheet_combo.currentText()
        parcel = self.parcel_edit.text()
        meta = {"vdc": vdc, "ward": ward, "sheet": sheet, "parcel": parcel}
        if not (vdc and parcel):
            QMessageBox.warning(self, "Error", "Please select all fields and enter a parcel number.")
            return
        found = False
        if ward == "(No Sheet)" or not ward:
            vdc_path = os.path.join(self.folder, vdc)
            for img in os.listdir(vdc_path):
                m = re.match(r"(\d+)-(\d+)\.jpe?g", img, re.IGNORECASE)
                if m and int(m.group(1)) <= int(parcel) <= int(m.group(2)):
                    image_path = os.path.join(vdc_path, img)
                    found = True
                    viewer = ImageViewerWindow(image_path, config=self.config, meta=meta, doc_type=self.doc_type)
                    viewer.show()
                    viewer.raise_()
                    viewer.activateWindow()
                    self._last_viewer = viewer
                    break
        else:
            sheet_path = os.path.join(self.folder, vdc, ward, sheet)
            for img in os.listdir(sheet_path):
                m = re.match(r"(\d+)-(\d+)\.jpe?g", img, re.IGNORECASE)
                if m and int(m.group(1)) <= int(parcel) <= int(m.group(2)):
                    image_path = os.path.join(sheet_path, img)
                    found = True
                    viewer = ImageViewerWindow(image_path, config=self.config, meta=meta, doc_type=self.doc_type)
                    viewer.show()
                    viewer.raise_()
                    viewer.activateWindow()
                    self._last_viewer = viewer
                    break
        if not found:
            QMessageBox.warning(self, "Not Found", "Parcel not found in this location.")

    def finalize_doc(self):
        doc_mgr = self.get_doc_mgr()
        if not doc_mgr.is_loaded():
            QMessageBox.information(self, "No Document", f"There is no active {self.doc_type.title()} to save.")
            return
        dlg = FieldbookBottomTextDialog(self)
        if dlg.exec_() == QDialog.Accepted:
            info = dlg.get_values()
            doc_mgr.footer_info = info
        else:
            return
        save_path, _ = QFileDialog.getSaveFileName(self, f"Save {self.doc_type.title()}", "", "Word Files (*.docx)")
        if save_path:
            doc_mgr.save(save_path)
            QMessageBox.information(self, "Saved", f"Document saved: {save_path}\nDocument cleared.")
            doc_mgr.close()

    def print_doc(self):
        import platform
        doc_mgr = self.get_doc_mgr()
        if not doc_mgr.is_loaded():
            QMessageBox.information(self, "No Document", f"There is no active {self.doc_type.title()} to print.")
            return
        if not getattr(doc_mgr, "footer_info", None):
            dlg = FieldbookBottomTextDialog(self)
            if dlg.exec_() == QDialog.Accepted:
                info = dlg.get_values()
                doc_mgr.footer_info = info
            else:
                return
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
            temp_docx_path = tf.name
            doc_mgr.save(temp_docx_path)
        try:
            if platform.system() == "Darwin":
                os.system(f'open "{temp_docx_path}"')
            elif os.name == "nt":
                os.startfile(temp_docx_path)
            elif platform.system().startswith("linux"):
                os.system(f'xdg-open "{temp_docx_path}"')
            else:
                raise OSError("Unsupported OS for auto-open")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open DOCX for preview: {e}")

class MainWindow(QMainWindow):
    def __init__(self, db, config):
        super().__init__()
        self.db = db
        self.config = config
        self.setWindowTitle("Fieldbook and Plot Register Viewer")
        self.resize(1200, 800)
        self.username = None
        self.role = None
        self.stacked = QStackedWidget()
        self.setCentralWidget(self.stacked)
        self.init_menu()
        self.show_login()

    def init_menu(self):
        menubar = self.menuBar()

        # File Menu
        self.menu_file = menubar.addMenu("File")
        self.action_print_fieldbook = QAction("Print Fieldbook", self)
        self.action_print_fieldbook.triggered.connect(self.print_fieldbook)
        self.menu_file.addAction(self.action_print_fieldbook)

        self.action_print_plotregister = QAction("Print Plot Register", self)
        self.action_print_plotregister.triggered.connect(self.print_plotregister)
        self.menu_file.addAction(self.action_print_plotregister)

        self.menu_file.addSeparator()

        self.action_logout = QAction("Logout", self)
        self.action_logout.triggered.connect(self.logout)
        self.menu_file.addAction(self.action_logout)
        self.menu_file.setEnabled(False)

        # Setup Menu
        self.menu_setup = menubar.addMenu("Setup")
        self.action_set_fieldbook = QAction("Set Fieldbook Folder", self)
        self.action_set_fieldbook.triggered.connect(self.set_fieldbook_folder)
        self.menu_setup.addAction(self.action_set_fieldbook)

        self.action_set_plotregister = QAction("Set Plot Register Folder", self)
        self.action_set_plotregister.triggered.connect(self.set_plotregister_folder)
        self.menu_setup.addAction(self.action_set_plotregister)

        self.menu_setup.addSeparator()

        self.action_load_template = QAction("Load Fieldbook Template", self)
        self.action_load_template.triggered.connect(self.load_fieldbook_template)
        self.menu_setup.addAction(self.action_load_template)

        self.action_load_plotregister_template = QAction("Load Plot Register Template", self)
        self.action_load_plotregister_template.triggered.connect(self.load_plotregister_template)
        self.menu_setup.addAction(self.action_load_plotregister_template)

        # About Menu
        self.menu_about = menubar.addMenu("About")
        self.action_app_info = QAction("Application Info", self)
        self.action_app_info.triggered.connect(self.show_app_info)
        self.menu_about.addAction(self.action_app_info)

        self.action_version = QAction("Version", self)
        self.action_version.triggered.connect(self.show_version)
        self.menu_about.addAction(self.action_version)

        self.action_manual = QAction("Manual", self)
        self.action_manual.triggered.connect(self.show_manual)
        self.menu_about.addAction(self.action_manual)

    def show_app_info(self):
        """Show application information."""
        QMessageBox.information(self, "Application Info",
            "Fieldbook and Plot Register Viewer\n\n"
            "An application for viewing and printing survey fieldbook and plot register images into the user defined templates. This application is developed by Er. Dinesh Bishwakarma who is a Survey Officer and working as a Survey Officer at Survey Office Kapilvastu, Nepal. \n\nDinesh had a thought to solve the issue of printing the fieldbook and plot register copies in the office. The major issues are the lack of indexing the registers in the digital environment, maintaining the common template within the office, and easy printing of those prints. Dinesh has developed this application not to solve the issue of the Survey Office Kapilvastu, but whole Survey Offices of Nepal. \n\n"
            "If you like this application, then please give feedback to dineshbishwakarma98@gmail.com . Your feedback and appreciation are necessary for the developer's motivation."
        )

    def show_version(self):
        """Show application version."""
        # Optionally, keep version as a class/global variable
        version = "v1.0.0"
        QMessageBox.information(self, "Version", f"Fieldbook and Plot Register Viewer r\nVersion: {version}")

    def show_manual(self):
        """Show application manual/usage instructions."""
        manual_text = (
            "<b>Fieldbook and Plot Register Viewer - User Manual</b><br><br>"
            "<b>Fieldbook Viewer:</b><br>"
            "- Set Fieldbook Folder and Template in Setup menu.<br>"
            "- Browse by VDC/Ward/Sheet/Parcel, open images and copy/crop/paste to fieldbook.<br>"
            "- Use 'Finalize' to save document, 'Print' to send to printer.<br><br>"
            "<b>Plot Register Viewer:</b><br>"
            "- Similar usage as Fieldbook, for plot register images.<br><br>"
            "<b>Other:</b><br>"
            "- Use the File menu for printing and logout.<br>"
            "- Use Setup for folder/template configuration.<br>"
            "- For any issue, contact dineshbishwakarma98@gmail.com <br><br>"
            "- Video Tutorial available here -> "
        )
        dlg = QDialog(self)
        dlg.setWindowTitle("User Manual")
        layout = QVBoxLayout(dlg)
        label = QLabel(manual_text)
        label.setWordWrap(True)
        label.setTextFormat(Qt.RichText)
        layout.addWidget(label)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dlg.accept)
        layout.addWidget(close_btn)
        dlg.exec_()

    def load_fieldbook_template(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Fieldbook Template", "", "Word Files (*.docx)")
        if file_path:
            self.config.set_folder("fieldbook_template", file_path)
            QMessageBox.information(self, "Template Loaded", "Fieldbook template loaded successfully.")

    def load_plotregister_template(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Plot Register Template", "", "Word Files (*.docx)")
        if file_path:
            self.config.set_folder("plotregister_template", file_path)
            QMessageBox.information(self, "Template Loaded", "Plot Register template loaded successfully.")

    def show_login(self):
        self.menu_file.setEnabled(False)
        while self.stacked.count() > 0:
            widget = self.stacked.widget(0)
            self.stacked.removeWidget(widget)
            widget.deleteLater()
        self.login_widget = LoginWidget(self.db, self.on_login)
        self.stacked.addWidget(self.login_widget)
        self.stacked.setCurrentWidget(self.login_widget)

    def on_login(self, username, role):
        self.username = username
        self.role = role
        self.menu_file.setEnabled(True)
        self.show_home()

    def show_home(self):
        home = QWidget()
        layout = QVBoxLayout(home)
        label = QLabel(f"Welcome, {self.username}!")
        label.setAlignment(Qt.AlignCenter)
        label.setStyleSheet("font-size: 22px; font-weight: bold; margin: 20px;")
        layout.addWidget(label)
        card_layout = QHBoxLayout()
        btn_fieldbook = QPushButton(QIcon.fromTheme("folder"), "Fieldbook Viewer")
        btn_fieldbook.setMinimumSize(220, 120)
        btn_fieldbook.setStyleSheet("font-size: 20px; border-radius: 16px;")
        btn_fieldbook.clicked.connect(self.show_fieldbook)
        btn_plotregister = QPushButton(QIcon.fromTheme("folder"), "Plot Register Viewer")
        btn_plotregister.setMinimumSize(220, 120)
        btn_plotregister.setStyleSheet("font-size: 20px; border-radius: 16px;")
        btn_plotregister.clicked.connect(self.show_plotregister)
        card_layout.addWidget(btn_fieldbook)
        card_layout.addWidget(btn_plotregister)
        layout.addLayout(card_layout)
        layout.addStretch()
        self.stacked.addWidget(home)
        self.stacked.setCurrentWidget(home)

    def show_fieldbook(self):
        folder = self.config.get_folder("fieldbook_folder")
        if not folder or not os.path.isdir(folder):
            QMessageBox.information(self, "Set Folder", "Please set the Fieldbook folder from the File menu.")
            return
        for idx in reversed(range(self.stacked.count())):
            widget = self.stacked.widget(idx)
            if isinstance(widget, BookViewer):
                self.stacked.removeWidget(widget)
                widget.deleteLater()
        def on_back():
            QApplication.clipboard().clear(mode=QClipboard.Clipboard)
            QApplication.clipboard().clear(mode=QClipboard.Selection)
            if fieldbook_doc_mgr.is_loaded():
                fieldbook_doc_mgr.close()
            self.show_home()
        self.fieldbook_viewer = BookViewer(self.config, "fieldbook_folder", "Fieldbook Viewer", doc_type="fieldbook", on_back=on_back)
        self.stacked.addWidget(self.fieldbook_viewer)
        self.stacked.setCurrentWidget(self.fieldbook_viewer)

    def show_plotregister(self):
        folder = self.config.get_folder("plotregister_folder")
        if not folder or not os.path.isdir(folder):
            QMessageBox.information(self, "Set Folder", "Please set the Plot Register folder from the File menu.")
            return
        for idx in reversed(range(self.stacked.count())):
            widget = self.stacked.widget(idx)
            if isinstance(widget, BookViewer):
                self.stacked.removeWidget(widget)
                widget.deleteLater()
        def on_back():
            QApplication.clipboard().clear(mode=QClipboard.Clipboard)
            QApplication.clipboard().clear(mode=QClipboard.Selection)
            if plotregister_doc_mgr.is_loaded():
                plotregister_doc_mgr.close()
            self.show_home()
        self.plotregister_viewer = BookViewer(self.config, "plotregister_folder", "Plot Register Viewer", doc_type="plotregister", on_back=on_back)
        self.stacked.addWidget(self.plotregister_viewer)
        self.stacked.setCurrentWidget(self.plotregister_viewer)

    def set_fieldbook_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Fieldbook Root Directory", os.getcwd())
        if folder:
            self.config.set_folder("fieldbook_folder", folder)
            QMessageBox.information(self, "Folder Set", "Fieldbook folder set successfully.")

    def set_plotregister_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Plot Register Root Directory", os.getcwd())
        if folder:
            self.config.set_folder("plotregister_folder", folder)
            QMessageBox.information(self, "Folder Set", "Plot Register folder set successfully.")

    def print_fieldbook(self):
        import platform
        if not fieldbook_doc_mgr.is_loaded():
            QMessageBox.information(self, "No Fieldbook", "Please finalize & save the fieldbook document first (use 'Save' in Fieldbook viewer).")
            return
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
            temp_path = tf.name
            fieldbook_doc_mgr.save(temp_path)
        try:
            if platform.system() == "Windows":
                os.startfile(temp_path, "print")
            elif platform.system() == "Darwin":
                subprocess.run(["open", "-a", "Microsoft Word", temp_path])
            else:
                subprocess.run(["libreoffice", "--pt", temp_path])
            QMessageBox.information(self, "Print", "Print dialog has been opened in your system's Word processor.\nPlease print from there.")
        except Exception as e:
            QMessageBox.warning(self, "Print Error", f"Could not open print dialog automatically.\nError: {str(e)}\nYou can open and print the saved DOCX file manually.")

    def print_plotregister(self):
        import platform
        if not plotregister_doc_mgr.is_loaded():
            QMessageBox.information(self, "No Plot Register", "Please finalize & save the Plot Register document first (use 'Save' in Plot Register viewer).")
            return
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
            temp_path = tf.name
            plotregister_doc_mgr.save(temp_path)
        try:
            if platform.system() == "Windows":
                os.startfile(temp_path, "print")
            elif platform.system() == "Darwin":
                subprocess.run(["open", "-a", "Microsoft Word", temp_path])
            else:
                subprocess.run(["libreoffice", "--pt", temp_path])
            QMessageBox.information(self, "Print", "Print dialog has been opened in your system's Word processor.\nPlease print from there.")
        except Exception as e:
            QMessageBox.warning(self, "Print Error", f"Could not open print dialog automatically.\nError: {str(e)}\nYou can open and print the saved DOCX file manually.")

    def logout(self):
        self.username = None
        self.role = None
        while self.stacked.count() > 0:
            widget = self.stacked.widget(0)
            self.stacked.removeWidget(widget)
            widget.deleteLater()
        self.show_login()

def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    palette = app.palette()
    palette.setColor(QPalette.Window, Qt.white)
    palette.setColor(QPalette.WindowText, Qt.black)
    palette.setColor(QPalette.Base, Qt.white)
    palette.setColor(QPalette.AlternateBase, Qt.lightGray)
    palette.setColor(QPalette.ToolTipBase, Qt.white)
    palette.setColor(QPalette.ToolTipText, Qt.black)
    palette.setColor(QPalette.Text, Qt.black)
    palette.setColor(QPalette.Button, Qt.white)
    palette.setColor(QPalette.ButtonText, Qt.black)
    palette.setColor(QPalette.Highlight, Qt.blue)
    palette.setColor(QPalette.HighlightedText, Qt.white)
    app.setPalette(palette)
    app.setStyleSheet("""
        QWidget {
            font-family: 'Segoe UI', 'Kalimati', 'Arial', sans-serif;
            font-size: 15px;
        }
        QMainWindow {
            background: #f7f7fa;
        }
        QGroupBox, QFrame {
            border: 1px solid #d0d0d0;
            border-radius: 12px;
            background: #ffffff;
            margin-top: 10px;
            padding: 12px;
        }
        QLabel {
            font-weight: 500;
        }
        QLineEdit, QComboBox, QTextEdit {
            border: 1.5px solid #b0b0b0;
            border-radius: 8px;
            padding: 6px 10px;
            background: #f9f9fc;
        }
        QPushButton {
            background-color: #1976d2;
            color: white;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 16px;
            font-weight: 600;
            margin: 6px 0;
        }
        QPushButton:hover {
            background-color: #1565c0;
        }
        QListWidget, QGraphicsView {
            border: 1.5px solid #b0b0b0;
            border-radius: 8px;
            background: #f9f9fc;
        }
        QHeaderView::section {
            background-color: #e3eafc;
            border: none;
            padding: 6px;
        }
    """)

    APPDATA = get_appdata_folder()
    os.makedirs(APPDATA, exist_ok=True)  # Ensure it exists!
    CONFIG_PATH = os.path.join(APPDATA, "config.json")
    DB_PATH = os.path.join(APPDATA, "users.db")

    db = UserDB(db_path=DB_PATH)
    config = Config(path=CONFIG_PATH)

    window = MainWindow(db, config)
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
